import anthropic
import json
import re
import os
import base64

client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))


def extract_text(file_path: str, filename: str) -> str:
    """ファイルからテキストを抽出"""
    ext = filename.lower().rsplit(".", 1)[-1]

    if ext == "pdf":
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                pages = []
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
                    for table in page.extract_tables():
                        for row in table:
                            pages.append("\t".join(str(c) if c else "" for c in row))
                return "\n".join(pages)
        except Exception as e:
            return f"[PDFエラー: {e}]"

    elif ext in ("xlsx", "xls"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            lines = []
            for ws in wb.worksheets:
                lines.append(f"=== {ws.title} ===")
                for row in ws.iter_rows(values_only=True):
                    if any(c is not None for c in row):
                        lines.append("\t".join(str(c) if c is not None else "" for c in row))
            return "\n".join(lines)
        except Exception as e:
            return f"[Excelエラー: {e}]"

    elif ext == "docx":
        try:
            from docx import Document
            doc = Document(file_path)
            lines = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    lines.append("\t".join(c.text for c in row.cells))
            return "\n".join(lines)
        except Exception as e:
            return f"[Wordエラー: {e}]"

    elif ext in ("jpg", "jpeg", "png", "webp"):
        try:
            with open(file_path, "rb") as f:
                image_data = base64.standard_b64encode(f.read()).decode("utf-8")
            mt = {"jpg": "image/jpeg", "jpeg": "image/jpeg",
                  "png": "image/png", "webp": "image/webp"}.get(ext, "image/jpeg")
            resp = client.messages.create(
                model="claude-opus-4-5-20251101",
                max_tokens=2000,
                messages=[{"role": "user", "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": mt, "data": image_data}},
                    {"type": "text", "text": "この書類画像のテキストをすべて書き起こしてください。品目名・数量・単価・金額を含むすべての文字を読み取ってください。"}
                ]}]
            )
            return resp.content[0].text
        except Exception as e:
            return f"[画像エラー: {e}]"

    return ""


def detect_doc_type(text: str, filename: str) -> str:
    """書類の種類を自動判定"""
    combined = (text[:500] + filename).lower()
    if any(k in combined for k in ["見積", "estimate", "御見積"]):
        return "見積書"
    if any(k in combined for k in ["請求", "invoice", "インボイス"]):
        return "請求書"
    if any(k in combined for k in ["納品", "delivery", "納入"]):
        return "納品書"
    if any(k in combined for k in ["発注", "purchase", "注文"]):
        return "発注書"
    if any(k in combined for k in ["領収", "receipt", "レシート"]):
        return "領収書"
    return "見積書"


def extract_with_claude(text: str, filename: str, doc_type: str) -> dict:
    """Claude AIで書類情報を抽出"""
    prompt = f"""以下は「{filename}」という{doc_type}から抽出したテキストです。
情報を正確に抽出してJSON形式のみで返してください。説明文や```は不要です。

テキスト:
{text[:5000]}

返答するJSONの形式:
{{
  "doc_type": "{doc_type}",
  "doc_number": "書類番号（見積番号・請求書番号など。なければ空文字）",
  "client_name": "取引先・顧客・発注元の名前（なければ空文字）",
  "doc_date": "書類の日付（YYYY-MM-DD形式、なければ空文字）",
  "due_date": "支払期限または納期（YYYY-MM-DD形式、なければ空文字）",
  "currency": "JPY",
  "items": [
    {{
      "name": "品目名・商品名・サービス名",
      "quantity": 数量（数値、不明なら1）,
      "unit": "単位（個/式/時間/月 など）",
      "unit_price": 単価（数値）,
      "total_price": 合計（数値）,
      "note": "備考・仕様"
    }}
  ],
  "total_amount": 合計金額（税込or税抜、見つかった方の数値）,
  "tax_amount": 消費税額（数値、なければ0）,
  "memo": "特記事項・備考（支払条件・納期条件など）"
}}"""

    resp = client.messages.create(
        model="claude-opus-4-5-20251101",
        max_tokens=3000,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = resp.content[0].text.strip()
    raw = re.sub(r"```(?:json)?|```", "", raw).strip()
    m = re.search(r'\{.*\}', raw, re.DOTALL)
    if m:
        raw = m.group(0)
    return json.loads(raw)


def search_market_price(item_name: str) -> dict:
    """品目の相場価格をClaude AIで推定"""
    prompt = f"""「{item_name}」の日本市場での一般的な相場価格を調査して教えてください。

以下のJSON形式のみで返答してください（説明文・```不要）:
{{
  "market_low": 相場の下限価格（数値・円、不明なら0）,
  "market_high": 相場の上限価格（数値・円、不明なら0）,
  "unit": "価格の単位（1個あたり・1時間あたり・1式など）",
  "memo": "相場の根拠・注意事項・地域差など（100文字以内）",
  "reliability": "high/medium/low（推定の信頼度）"
}}"""

    resp = client.messages.create(
        model="claude-opus-4-5-20251101",
        max_tokens=500,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = resp.content[0].text.strip()
    raw = re.sub(r"```(?:json)?|```", "", raw).strip()
    m = re.search(r'\{.*\}', raw, re.DOTALL)
    if m:
        raw = m.group(0)
    return json.loads(raw)


def generate_estimate(client_name: str, items: list, company_name: str, memo: str) -> dict:
    """見積書データを自動生成"""
    items_text = "\n".join(
        f"- {it['name']}: 数量{it.get('quantity',1)}{it.get('unit','')}, 単価{it.get('unit_price',0)}円"
        for it in items
    )
    prompt = f"""以下の情報から見積書データを生成してください。

発行会社: {company_name}
取引先: {client_name}
品目:
{items_text}
備考: {memo}

以下のJSON形式のみで返答してください:
{{
  "doc_number": "見積番号（EST-YYYYMMDDの形式）",
  "doc_date": "今日の日付（YYYY-MM-DD）",
  "due_date": "有効期限（30日後のYYYY-MM-DD）",
  "items": [
    {{
      "name": "品目名",
      "quantity": 数量,
      "unit": "単位",
      "unit_price": 単価（数値）,
      "total_price": 小計（数値）,
      "note": "仕様・備考"
    }}
  ],
  "subtotal": 小計合計（数値）,
  "tax_rate": 0.1,
  "tax_amount": 消費税（数値）,
  "total_amount": 税込合計（数値）,
  "memo": "お支払い条件・備考など"
}}"""

    resp = client.messages.create(
        model="claude-opus-4-5-20251101",
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = resp.content[0].text.strip()
    raw = re.sub(r"```(?:json)?|```", "", raw).strip()
    m = re.search(r'\{.*\}', raw, re.DOTALL)
    if m:
        raw = m.group(0)
    return json.loads(raw)
